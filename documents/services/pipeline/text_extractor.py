from __future__ import annotations
from dataclasses import dataclass
import csv
import io

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class ExtractResult:
    text: str
    word_count: int
    char_count: int


def _count_words(text: str) -> int:
    # นับแบบง่าย: split ด้วย whitespace (เหมาะกับ prototype)
    return len([w for w in text.split() if w.strip()])


def extract_text_bytes(file_bytes: bytes, file_ext: str) -> ExtractResult:
    ext = (file_ext or "").lower().lstrip(".")

    if ext == "txt":
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("utf-8", errors="ignore")

    elif ext == "csv":
        text = _extract_csv_bytes(file_bytes)

    elif ext == "pdf":
        text = _extract_pdf_bytes(file_bytes)

    elif ext == "docx":
        text = _extract_docx_bytes(file_bytes)

    else:
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    text = (text or "").strip()
    return ExtractResult(text=text, word_count=_count_words(text), char_count=len(text))


def _extract_pdf_bytes(b: bytes) -> str:
    reader = PdfReader(io.BytesIO(b))
    parts = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx_bytes(b: bytes) -> str:
    doc = DocxDocument(io.BytesIO(b))
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts)


def _extract_csv_bytes(b: bytes) -> str:
    try:
        s = b.decode("utf-8")
    except UnicodeDecodeError:
        s = b.decode("utf-8", errors="ignore")

    f = io.StringIO(s)
    reader = csv.reader(f)
    lines = []
    for row in reader:
        row_clean = [c.strip() for c in row if c and c.strip()]
        if row_clean:
            lines.append(", ".join(row_clean))
    return "\n".join(lines)