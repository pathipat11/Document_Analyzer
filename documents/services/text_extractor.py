from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
import csv
import io

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class ExtractResult:
    text: str
    word_count: int
    char_count: int


def _count_words(text: str) -> int:
    # นับแบบง่าย: split ด้วย whitespace (เหมาะกับ prototype)
    return len([w for w in text.split() if w.strip()])


def extract_text(file_path: str, file_ext: str) -> ExtractResult:
    ext = (file_ext or "").lower().lstrip(".")
    path = Path(file_path)

    if ext == "txt":
        text = path.read_text(encoding="utf-8", errors="ignore")

    elif ext == "csv":
        # แปลง csv เป็นข้อความอ่านง่าย (join ทุก cell)
        text = _extract_csv(path)

    elif ext == "pdf":
        text = _extract_pdf(path)

    elif ext == "docx":
        text = _extract_docx(path)

    else:
        # future: รองรับ text-readable อื่น ๆ
        text = path.read_text(encoding="utf-8", errors="ignore")

    # normalize ข้อความเล็กน้อย
    text = (text or "").strip()

    return ExtractResult(
        text=text,
        word_count=_count_words(text),
        char_count=len(text),
    )


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        t = t.strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tables (เผื่อเอกสารมีตาราง)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    # พยายามอ่านแบบ utf-8 ก่อน ถ้าไม่ได้ค่อย fallback ignore
    raw = path.read_bytes()
    try:
        s = raw.decode("utf-8")
    except UnicodeDecodeError:
        s = raw.decode("utf-8", errors="ignore")

    f = io.StringIO(s)
    reader = csv.reader(f)
    lines: list[str] = []
    for row in reader:
        row_clean = [c.strip() for c in row if c and c.strip()]
        if row_clean:
            lines.append(", ".join(row_clean))
    return "\n".join(lines)
